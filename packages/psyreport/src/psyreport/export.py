"""Export psy results to LaTeX and Word (.docx), reusing the APA text renderers."""
from __future__ import annotations

import pandas as pd

from .report import report


def _title_and_note(text):
    """Split an APA text block into (title lines, table?, note)."""
    return text


def to_latex(result, caption=None, label=None):
    """Return a LaTeX table for a result that carries a table.

    Falls back to a LaTeX comment with the APA sentence for inline-only results
    (t-test, correlations text, etc.).
    """
    table = getattr(result, "table", None)
    apa = report(result)
    if table is None:
        return "% " + apa.replace("\n", "\n% ")
    cap = caption or _default_caption(result)
    body = table.to_latex(index=True, escape=True)
    lab = f"\\label{{{label}}}\n" if label else ""
    return (f"\\begin{{table}}[htbp]\n\\centering\n"
            f"\\caption{{{cap}}}\n{lab}{body}\\end{{table}}")


def to_docx(result, path, caption=None):
    """Write an APA-style .docx with the result's table and note. Returns path.

    Requires python-docx (imported lazily).
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    apa = report(result)
    lines = apa.split("\n")
    doc.add_paragraph(caption or _default_caption(result))

    table = getattr(result, "table", None)
    if table is not None:
        df = table.reset_index() if table.index.name or not _is_default_index(table) else table
        t = doc.add_table(rows=1, cols=len(df.columns))
        t.style = "Light Grid Accent 1"
        for j, col in enumerate(df.columns):
            t.rows[0].cells[j].text = str(col)
        for _, row in df.iterrows():
            cells = t.add_row().cells
            for j, val in enumerate(row):
                cells[j].text = _fmt(val)
        # APA note = the "Note." line, if present
        note = next((ln for ln in lines if ln.startswith("Note.")), "")
        if note:
            doc.add_paragraph(note)
    else:
        for ln in lines:
            doc.add_paragraph(ln)

    doc.save(path)
    return path


def _is_default_index(df):
    return isinstance(df.index, pd.RangeIndex)


def _fmt(val):
    if isinstance(val, float):
        return f"{val:.3f}"
    return str(val)


def _default_caption(result):
    kind = getattr(result, "kind", "result")
    titles = {
        "efa": "Exploratory Factor Analysis: Factor Loadings",
        "cfa": "Confirmatory Factor Analysis: Standardized Loadings",
        "corr_matrix": "Correlations",
        "table1": "Sample Characteristics",
        "describe": "Descriptive Statistics",
        "freq": "Frequency Distribution",
        "alpha": "Item Statistics and Reliability",
        "linreg": "Linear Regression Coefficients",
        "logreg": "Logistic Regression Coefficients",
    }
    return titles.get(kind, "Results")
