"""Export to LaTeX and Word, plus EFA/CFA APA renderers."""
import numpy as np
import pandas as pd
import pytest
from docx import Document

from psystats import corr_matrix, describe
from psymetrics import efa, cfa, alpha
from psyreport import report, to_latex, to_docx


@pytest.fixture
def items():
    rng = np.random.default_rng(7)
    n = 250
    f1, f2 = rng.normal(size=n), rng.normal(size=n)
    d = {f"a{i}": f1 + rng.normal(scale=.6, size=n) for i in range(1, 4)}
    d.update({f"b{i}": f2 + rng.normal(scale=.6, size=n) for i in range(1, 4)})
    return pd.DataFrame(d)


def test_efa_apa_render(items):
    s = report(efa(items, 2, rotation="promax"))
    assert "Exploratory factor analysis" in s and "h²" in s and "Note." in s


def test_cfa_apa_render():
    pytest.importorskip("semopy")  # optional SEM engine
    from semopy.examples import holzinger39
    s = report(cfa(holzinger39.get_model(), holzinger39.get_data()))
    assert "χ²(" in s and "CFI =" in s and "RMSEA =" in s


def test_to_latex_table(items):
    tex = to_latex(corr_matrix(items), caption="Correlations", label="tab:cor")
    assert "\\begin{table}" in tex and "\\caption{Correlations}" in tex
    assert "\\label{tab:cor}" in tex and "\\begin{tabular}" in tex


def test_to_latex_inline_fallback(items):
    # alpha carries a table, so it produces a real table; use a text-only case
    from psystats import ttest
    df = pd.DataFrame({"y": np.r_[np.zeros(20), np.ones(20)] + np.random.default_rng(1).normal(size=40),
                       "g": ["a"] * 20 + ["b"] * 20})
    tex = to_latex(ttest(df, "y", "g"))
    assert tex.startswith("%")  # commented APA sentence, no table


def test_to_docx_creates_readable_file(items, tmp_path):
    out = tmp_path / "efa.docx"
    to_docx(efa(items, 2), str(out))
    doc = Document(str(out))
    assert len(doc.tables) == 1
    # header row has the factor columns plus communality
    header = [c.text for c in doc.tables[0].rows[0].cells]
    assert any("F1" in h for h in header)


def test_to_docx_describe(items, tmp_path):
    out = tmp_path / "desc.docx"
    to_docx(describe(items), str(out))
    assert out.exists() and out.stat().st_size > 0
